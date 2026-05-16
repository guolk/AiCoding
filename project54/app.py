from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, timedelta
import os
import json
from docx import Document
from docx.shared import Inches
import markdown
import io

app = Flask(__name__)
CORS(app)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///meeting_system.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
db = SQLAlchemy(app)


class Participant(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class Meeting(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    date = db.Column(db.DateTime, nullable=False)
    agenda = db.Column(db.Text)
    content = db.Column(db.Text)
    summary = db.Column(db.Text)
    audio_transcript = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    participants = db.relationship('MeetingParticipant', backref='meeting', cascade='all, delete-orphan')
    topics = db.relationship('Topic', backref='meeting', cascade='all, delete-orphan')
    decisions = db.relationship('Decision', backref='meeting', cascade='all, delete-orphan')
    todos = db.relationship('Todo', backref='meeting', cascade='all, delete-orphan')


class MeetingParticipant(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    meeting_id = db.Column(db.Integer, db.ForeignKey('meeting.id'), nullable=False)
    participant_id = db.Column(db.Integer, db.ForeignKey('participant.id'), nullable=False)
    participant = db.relationship('Participant')


class Topic(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    meeting_id = db.Column(db.Integer, db.ForeignKey('meeting.id'), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    content = db.Column(db.Text)
    order = db.Column(db.Integer, default=0)


class Decision(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    meeting_id = db.Column(db.Integer, db.ForeignKey('meeting.id'), nullable=False)
    topic_id = db.Column(db.Integer, db.ForeignKey('topic.id'))
    content = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class Todo(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    meeting_id = db.Column(db.Integer, db.ForeignKey('meeting.id'), nullable=False)
    topic_id = db.Column(db.Integer, db.ForeignKey('topic.id'))
    content = db.Column(db.Text, nullable=False)
    assignee_id = db.Column(db.Integer, db.ForeignKey('participant.id'))
    assignee = db.relationship('Participant')
    due_date = db.Column(db.DateTime)
    status = db.Column(db.String(20), default='pending')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    completed_at = db.Column(db.DateTime)


class Risk(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    meeting_id = db.Column(db.Integer, db.ForeignKey('meeting.id'), nullable=False)
    content = db.Column(db.Text, nullable=False)
    level = db.Column(db.String(20))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


with app.app_context():
    db.create_all()


@app.route('/api/participants', methods=['GET', 'POST'])
def manage_participants():
    if request.method == 'POST':
        data = request.json
        participant = Participant(name=data['name'], email=data.get('email'))
        db.session.add(participant)
        db.session.commit()
        return jsonify({'id': participant.id, 'name': participant.name, 'email': participant.email})
    participants = Participant.query.all()
    return jsonify([{'id': p.id, 'name': p.name, 'email': p.email} for p in participants])


def parse_iso_datetime(date_str):
    """Python 3.6兼容的ISO日期解析"""
    try:
        if '.' in date_str:
            date_str = date_str.split('.')[0]
        return datetime.strptime(date_str, '%Y-%m-%dT%H:%M:%S')
    except:
        try:
            return datetime.strptime(date_str, '%Y-%m-%d %H:%M:%S')
        except:
            return datetime.strptime(date_str[:10], '%Y-%m-%d')


@app.route('/api/meetings', methods=['GET', 'POST'])
def manage_meetings():
    if request.method == 'POST':
        data = request.json
        meeting = Meeting(
            title=data['title'],
            date=parse_iso_datetime(data['date']),
            agenda=data.get('agenda', ''),
            content=data.get('content', '')
        )
        
        for p_id in data.get('participant_ids', []):
            mp = MeetingParticipant(participant_id=p_id)
            meeting.participants.append(mp)
        
        db.session.add(meeting)
        db.session.commit()
        return jsonify({'id': meeting.id, 'message': 'Meeting created successfully'})
    
    meetings = Meeting.query.order_by(Meeting.date.desc()).all()
    result = []
    for m in meetings:
        result.append({
            'id': m.id,
            'title': m.title,
            'date': m.date.isoformat(),
            'participants': [{'id': mp.participant.id, 'name': mp.participant.name} for mp in m.participants]
        })
    return jsonify(result)


@app.route('/api/meetings/<int:meeting_id>', methods=['GET', 'PUT', 'DELETE'])
def manage_meeting(meeting_id):
    meeting = Meeting.query.get_or_404(meeting_id)
    
    if request.method == 'GET':
        return jsonify({
            'id': meeting.id,
            'title': meeting.title,
            'date': meeting.date.isoformat(),
            'agenda': meeting.agenda,
            'content': meeting.content,
            'summary': meeting.summary,
            'audio_transcript': meeting.audio_transcript,
            'participants': [{'id': mp.participant.id, 'name': mp.participant.name, 'email': mp.participant.email} 
                           for mp in meeting.participants],
            'topics': [{'id': t.id, 'title': t.title, 'content': t.content, 'order': t.order} 
                      for t in sorted(meeting.topics, key=lambda x: x.order)],
            'decisions': [{'id': d.id, 'content': d.content} for d in meeting.decisions],
            'todos': [{'id': t.id, 'content': t.content, 'assignee': t.assignee.name if t.assignee else None,
                      'due_date': t.due_date.isoformat() if t.due_date else None, 'status': t.status} 
                     for t in meeting.todos]
        })
    
    elif request.method == 'PUT':
        data = request.json
        meeting.title = data.get('title', meeting.title)
        meeting.content = data.get('content', meeting.content)
        meeting.summary = data.get('summary', meeting.summary)
        meeting.audio_transcript = data.get('audio_transcript', meeting.audio_transcript)
        db.session.commit()
        return jsonify({'message': 'Meeting updated successfully'})
    
    elif request.method == 'DELETE':
        db.session.delete(meeting)
        db.session.commit()
        return jsonify({'message': 'Meeting deleted successfully'})


@app.route('/api/meetings/<int:meeting_id>/topics', methods=['POST'])
def add_topic(meeting_id):
    data = request.json
    topic = Topic(
        meeting_id=meeting_id,
        title=data['title'],
        content=data.get('content', ''),
        order=data.get('order', 0)
    )
    db.session.add(topic)
    db.session.commit()
    return jsonify({'id': topic.id})


@app.route('/api/meetings/<int:meeting_id>/todos', methods=['GET', 'POST'])
def manage_todos(meeting_id):
    if request.method == 'POST':
        data = request.json
        todo = Todo(
            meeting_id=meeting_id,
            content=data['content'],
            assignee_id=data.get('assignee_id'),
            due_date=parse_iso_datetime(data['due_date']) if data.get('due_date') else None
        )
        db.session.add(todo)
        db.session.commit()
        return jsonify({'id': todo.id})
    
    todos = Todo.query.filter_by(meeting_id=meeting_id).all()
    return jsonify([{
        'id': t.id,
        'content': t.content,
        'assignee': t.assignee.name if t.assignee else None,
        'due_date': t.due_date.isoformat() if t.due_date else None,
        'status': t.status
    } for t in todos])


@app.route('/api/todos/<int:todo_id>', methods=['PUT'])
def update_todo(todo_id):
    todo = Todo.query.get_or_404(todo_id)
    data = request.json
    todo.status = data.get('status', todo.status)
    if todo.status == 'completed' and not todo.completed_at:
        todo.completed_at = datetime.utcnow()
    db.session.commit()
    return jsonify({'message': 'Todo updated'})


@app.route('/api/todos', methods=['GET'])
def get_all_todos():
    status = request.args.get('status')
    query = Todo.query
    if status:
        query = query.filter_by(status=status)
    todos = query.order_by(Todo.created_at.desc()).all()
    return jsonify([{
        'id': t.id,
        'content': t.content,
        'assignee': t.assignee.name if t.assignee else None,
        'due_date': t.due_date.isoformat() if t.due_date else None,
        'status': t.status,
        'meeting_title': t.meeting.title,
        'meeting_date': t.meeting.date.isoformat()
    } for t in todos])


@app.route('/api/decisions', methods=['GET'])
def get_all_decisions():
    decisions = Decision.query.order_by(Decision.created_at.desc()).all()
    return jsonify([{
        'id': d.id,
        'content': d.content,
        'meeting_title': d.meeting.title,
        'meeting_date': d.meeting.date.isoformat()
    } for d in decisions])


@app.route('/api/meetings/<int:meeting_id>/decisions', methods=['POST'])
def add_decision(meeting_id):
    data = request.json
    decision = Decision(meeting_id=meeting_id, content=data['content'])
    db.session.add(decision)
    db.session.commit()
    return jsonify({'id': decision.id})


@app.route('/api/search', methods=['GET'])
def search_meetings():
    keyword = request.args.get('keyword', '')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    participant_id = request.args.get('participant_id')
    
    print(f"搜索参数: keyword={keyword}, start_date={start_date}, end_date={end_date}, participant_id={participant_id}")
    
    query = Meeting.query
    
    if keyword:
        query = query.filter(
            (Meeting.title.contains(keyword)) |
            (Meeting.content.contains(keyword)) |
            (Meeting.agenda.contains(keyword))
        )
    
    if start_date:
        start_dt = parse_iso_datetime(start_date + 'T00:00:00')
        query = query.filter(Meeting.date >= start_dt)
    if end_date:
        end_dt = parse_iso_datetime(end_date + 'T23:59:59')
        query = query.filter(Meeting.date <= end_dt)
    
    if participant_id and participant_id != '':
        query = query.join(MeetingParticipant).filter(
            MeetingParticipant.participant_id == int(participant_id)
        )
    
    meetings = query.order_by(Meeting.date.desc()).all()
    print(f"找到 {len(meetings)} 个会议")
    result = []
    for m in meetings:
        result.append({
            'id': m.id,
            'title': m.title,
            'date': m.date.strftime('%Y-%m-%dT%H:%M:%S')
        })
    return jsonify(result)


@app.route('/api/statistics', methods=['GET'])
def get_statistics():
    thirty_days_ago = datetime.utcnow() - timedelta(days=30)
    
    total_meetings = Meeting.query.count()
    recent_meetings = Meeting.query.filter(Meeting.date >= thirty_days_ago).count()
    
    participant_stats = db.session.query(
        Participant.name,
        db.func.count(MeetingParticipant.id).label('count')
    ).join(MeetingParticipant).group_by(Participant.id).order_by(db.desc('count')).limit(10).all()
    
    pending_todos = Todo.query.filter_by(status='pending').count()
    completed_todos = Todo.query.filter_by(status='completed').count()
    
    return jsonify({
        'total_meetings': total_meetings,
        'recent_meetings': recent_meetings,
        'pending_todos': pending_todos,
        'completed_todos': completed_todos,
        'top_participants': [{'name': p[0], 'count': p[1]} for p in participant_stats]
    })


@app.route('/api/meetings/<int:meeting_id>/export/markdown', methods=['GET'])
def export_markdown(meeting_id):
    meeting = Meeting.query.get_or_404(meeting_id)
    
    md_content = f"# {meeting.title}\n\n"
    md_content += f"**日期:** {meeting.date.strftime('%Y-%m-%d %H:%M')}\n\n"
    
    md_content += "## 参会人员\n"
    for mp in meeting.participants:
        md_content += f"- {mp.participant.name}\n"
    md_content += "\n"
    
    if meeting.agenda:
        md_content += "## 会议议程\n"
        md_content += f"{meeting.agenda}\n\n"
    
    for topic in sorted(meeting.topics, key=lambda x: x.order):
        md_content += f"## {topic.title}\n"
        if topic.content:
            md_content += f"{topic.content}\n\n"
        
        topic_decisions = [d for d in meeting.decisions if d.topic_id == topic.id]
        if topic_decisions:
            md_content += "### 决策\n"
            for d in topic_decisions:
                md_content += f"- [x] {d.content}\n"
            md_content += "\n"
        
        topic_todos = [t for t in meeting.todos if t.topic_id == topic.id]
        if topic_todos:
            md_content += "### 待办事项\n"
            for t in topic_todos:
                assignee = f" @{t.assignee.name}" if t.assignee else ""
                due = f" 截止: {t.due_date.strftime('%Y-%m-%d')}" if t.due_date else ""
                md_content += f"- [ ] {t.content}{assignee}{due}\n"
            md_content += "\n"
    
    if meeting.summary:
        md_content += "## 会议摘要\n"
        md_content += f"{meeting.summary}\n"
    
    output = io.BytesIO()
    output.write(md_content.encode('utf-8'))
    output.seek(0)
    
    return send_file(output, mimetype='text/markdown', 
                     as_attachment=True, 
                     download_name=f"{meeting.title}.md")


@app.route('/api/meetings/<int:meeting_id>/export/word', methods=['GET'])
def export_word(meeting_id):
    meeting = Meeting.query.get_or_404(meeting_id)
    
    doc = Document()
    doc.add_heading(meeting.title, 0)
    
    doc.add_paragraph(f"日期: {meeting.date.strftime('%Y-%m-%d %H:%M')}")
    
    doc.add_heading('参会人员', level=1)
    for mp in meeting.participants:
        doc.add_paragraph(f"{mp.participant.name}", style='List Bullet')
    
    if meeting.agenda:
        doc.add_heading('会议议程', level=1)
        doc.add_paragraph(meeting.agenda)
    
    for topic in sorted(meeting.topics, key=lambda x: x.order):
        doc.add_heading(topic.title, level=1)
        if topic.content:
            doc.add_paragraph(topic.content)
        
        topic_decisions = [d for d in meeting.decisions if d.topic_id == topic.id]
        if topic_decisions:
            doc.add_heading('决策', level=2)
            for d in topic_decisions:
                doc.add_paragraph(f"✓ {d.content}", style='List Bullet')
        
        topic_todos = [t for t in meeting.todos if t.topic_id == topic.id]
        if topic_todos:
            doc.add_heading('待办事项', level=2)
            for t in topic_todos:
                assignee = f" 负责人: {t.assignee.name}" if t.assignee else ""
                due = f" 截止: {t.due_date.strftime('%Y-%m-%d')}" if t.due_date else ""
                doc.add_paragraph(f"☐ {t.content}{assignee}{due}", style='List Bullet')
    
    if meeting.summary:
        doc.add_heading('会议摘要', level=1)
        doc.add_paragraph(meeting.summary)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(output, 
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True,
                     download_name=f"{meeting.title}.docx")


@app.route('/api/meetings/<int:meeting_id>/email-draft', methods=['GET'])
def get_email_draft(meeting_id):
    meeting = Meeting.query.get_or_404(meeting_id)
    
    recipients = [mp.participant.email for mp in meeting.participants if mp.participant.email]
    subject = f"会议纪要: {meeting.title}"
    
    body = f"各位参会者：\n\n"
    body += f"附件是「{meeting.title}」的会议纪要，请查收。\n\n"
    
    body += "待办事项汇总：\n"
    for todo in meeting.todos:
        if todo.status == 'pending':
            assignee = todo.assignee.name if todo.assignee else "未分配"
            body += f"- {todo.content} (负责人: {assignee})\n"
    
    body += "\n此致\n会议系统"
    
    return jsonify({
        'recipients': recipients,
        'subject': subject,
        'body': body
    })


@app.route('/')
def index():
    return send_file('index.html')


if __name__ == '__main__':
    app.run(debug=True, port=5000)
